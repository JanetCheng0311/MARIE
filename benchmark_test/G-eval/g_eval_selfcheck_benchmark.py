"""
g-eval style SelfCheck benchmark script

Notes (English):
- Reads JSON test items (or uses built-in `sample-1` example).
- If `selfcheckgpt` is installed, uses `SelfCheckMQAG` for sentence-level scores.
- If not installed, uses a simple `DummyChecker` so the script still runs.
- Uses `metadata.true` / `metadata.false` for a simple PASS/FAIL/AMBIGUOUS verdict.
- Outputs JSON results and attempts to write a .docx summary (falls back to .txt).

Spacy sentence-boundary note:
If you see the error:
    File "spacy/tokens/doc.pyx", line 926, in sents
    ValueError: [E030] Sentence boundaries unset. You can add the 'sentencizer' component to the pipeline with: `nlp.add_pipe('sentencizer')`.
Then install `spacy` and either add the `sentencizer` or a parser to your pipeline. This script will
automatically add `sentencizer` to the pipeline when loading spacy to avoid that error.

Usage example:
python3 benchmark_test/g_eval_selfcheck_benchmark.py --input path/to/items.json --output results/out.json

註記（繁體中文）:
- 讀入 JSON 格式的測試項目（若未提供 input 檔，會使用內建的 `sample-1` 範例）。
- 若已安裝 `selfcheckgpt`，程式會使用 `SelfCheckMQAG` 來對句子做評分；若未安裝，程式會使用
    一個簡易的 `DummyChecker` 當作備援，使腳本仍可執行。
- 判定規則：使用 `metadata.true` 與 `metadata.false` 做簡單的字串比對；
    - 只找到 true -> `PASS`
    - 只找到 false -> `FAIL`
    - 兩者都找到 -> `AMBIGUOUS`
    （如需不同判定行為，可改為只比對 `passage` 或加入中文分詞/模糊比對）
- Spacy 斷句錯誤（常見）:
    如果出現錯誤訊息：
        ValueError: [E030] Sentence boundaries unset.
    可用 `nlp.add_pipe('sentencizer')` 加入 sentencizer，或安裝 parser/senter 等元件。
    此腳本在載入 spacy 時會嘗試自動加入 `sentencizer`（若可用），以避免該錯誤。

此檔同時包含繁體中文註解以便閱讀。
"""

import argparse
import json
import os
from datetime import datetime
import numbers

# 嘗試載入 torch（用於判斷 cuda 是否可用），若無則降為 None
try:
    import torch
except Exception:
    torch = None

# 嘗試載入 spacy（斷句/切句），若無則降為 None，使用內建備援切句
try:
    import spacy
except Exception:
    spacy = None

# 嘗試載入 selfcheckgpt 的檢查器類別；若系統尚未安裝該套件，會使用內建 DummyChecker
try:
    from selfcheckgpt.modeling_selfcheck import SelfCheckMQAG
except Exception:
    SelfCheckMQAG = None


def evaluate_item(item, checker, nlp):
    """
    使用已安裝的斷句器(nlp)與 checker 對單一 item 做評分

    參數:
    - item: 單筆測試資料（包含 passage 與 sampled_passages）
    - checker: SelfCheckMQAG 或 DummyChecker（必須提供 predict()）
    - nlp: spacy 的 nlp 物件（若為 None，呼叫端需使用備援切句）

    回傳結構包含: id, passage, sentences, sampled_passages, scores
    """

    passage = item.get("passage", "")
    sampled = item.get("sampled_passages", [])

    # 使用 spacy 的句子分割（caller 已確保 nlp 非 None 時才呼叫此函式）
    sentences = [s.text.strip() for s in nlp(passage).sents]

    # 呼叫檢查器取得每句的分數
    scores = checker.predict(
        sentences=sentences,
        passage=passage,
        sampled_passages=sampled,
        num_questions_per_sent=item.get("num_questions_per_sent", 3),
        beta1=item.get("beta1", 0.8),
        beta2=item.get("beta2", 0.8),
    )

    return {
        "id": item.get("id"),
        "passage": passage,
        "sentences": sentences,
        "sampled_passages": sampled,
        "scores": scores,
    }


def evaluate_answer_verdict(item, passage, sentences, sampled_passages):
    """
    根據 item.metadata 裡的 true / false 標籤作簡單判定。

    規則（目前簡單字串比對）:
    - 若只找到 true 標籤 -> PASS
    - 若只找到 false 標籤 -> FAIL
    - 若同時找到 true 與 false -> AMBIGUOUS

    備註: 目前為字串包含比對，將來可改為只比對 passage 或使用中文分詞/模糊比對
    """

    meta = item.get("metadata", {})
    trues = [t.lower() for t in meta.get("true", [])]
    falses = [f.lower() for f in meta.get("false", [])]

    # 將 passage、sampled 與斷句結果合併為比對語料
    text_corpus = "\n".join([passage] + sampled_passages + sentences).lower()

    found_true = [t for t in trues if t and t in text_corpus]
    found_false = [f for f in falses if f and f in text_corpus]

    if found_true and not found_false:
        verdict = "PASS"
    elif found_true and found_false:
        verdict = "AMBIGUOUS"
    else:
        verdict = "FAIL"

    return {
        "found_true": found_true,
        "found_false": found_false,
        "verdict": verdict,
    }


def make_json_serializable(o):
    """Recursively convert common non-serializable objects to JSON-serializable types.

    Handles: numpy arrays/dtypes, torch tensors, sets, tuples, datetimes, and falls back to str().
    """
    # Primitive types
    if o is None or isinstance(o, (str, bool, numbers.Integral, numbers.Real)):
        return o

    # Containers
    if isinstance(o, (list, tuple, set)):
        return [make_json_serializable(x) for x in o]
    if isinstance(o, dict):
        return {str(k): make_json_serializable(v) for k, v in o.items()}

    # Lazy import numpy
    try:
        import numpy as _np

        if isinstance(o, _np.ndarray):
            return o.tolist()
        if isinstance(o, (_np.floating, _np.integer)):
            return o.item()
    except Exception:
        pass

    # torch tensors
    try:
        import torch as _torch

        if isinstance(o, _torch.Tensor):
            try:
                lst = o.cpu().detach().tolist()
                return lst
            except Exception:
                try:
                    return float(o.cpu().detach().item())
                except Exception:
                    return str(o)
    except Exception:
        pass

    # datetime-like
    try:
        if hasattr(o, 'isoformat'):
            return o.isoformat()
    except Exception:
        pass

    # Fallback: convert to string
    try:
        return str(o)
    except Exception:
        return None


def main():
    parser = argparse.ArgumentParser(description="Run SelfCheckMQAG benchmark (g-eval style)")
    parser.add_argument("--input", "-i", help="Input JSON file with items", default=None)
    parser.add_argument("--output", "-o", help="Output JSON file", default=None)
    args = parser.parse_args()

    # Load input items
    if args.input and os.path.exists(args.input):
        with open(args.input, "r", encoding="utf-8") as f:
            items = json.load(f)
    else:
        # Default sample items
        items = [
            {
                "id": "sample-1",
                "passage": "港鐵繼續提供八達通付款服務可以喺港鐵站、巴士同便利店付款，非常方便；天文台預測今個冬天氣溫偏低，但唔會落雪；而政府持續推動創新科技發展，探索太空相關研究計劃。",
                "sampled_passages": [
                    "聽講港鐵啱啱宣佈可以用八達通搭飛機。",
                    "天文台話今個冬天會落雪一星期咁耐。",
                    "政府計劃喺中環開放太空電梯，下個月試運行。",
                ],
            }
        ]

    # Determine device (fallback to cpu if torch missing)
    if torch is not None:
        device = "cuda" if torch.cuda.is_available() else "cpu"
    else:
        device = "cpu"

    # Instantiate checker; if SelfCheckMQAG not available, use a simple fallback
    if SelfCheckMQAG is not None:
        checker = SelfCheckMQAG(device=device)
    else:
        class DummyChecker:
            def __init__(self, **kwargs):
                pass

            def predict(self, sentences, passage, sampled_passages, **kwargs):
                # Simple heuristic: if any sentence contains words from item-level flags
                scores = []
                for sent in sentences:
                    s = sent.lower()
                    score = 0.5
                    # penalize if sampled passages contain obvious contradictions
                    for sp in sampled_passages:
                        if any(tok in sp for tok in ["長頸鹿", "狗", "生高", "生仔"]):
                            score = max(score, 0.9)
                    # reward if passage mentions likely-true keywords
                    if any(tok in s for tok in ["龜", "龜苓膏", "烏龜"]):
                        score = min(score, 0.1)
                    scores.append(score)
                return scores

        checker = DummyChecker()

    # Load a lightweight tokenizer/sentence splitter; fallback if spacy missing
    if spacy is not None:
        try:
            # Try to load a small English model first
            nlp = spacy.load("en_core_web_sm")
        except Exception:
            try:
                # Fall back to a blank English pipeline
                nlp = spacy.blank("en")
            except Exception:
                nlp = None

    # If we have a spacy pipeline but no sentencizer or parser, add the sentencizer
    # This prevents the ValueError about unset sentence boundaries when accessing `doc.sents`.
    # English note: If you see "Sentence boundaries unset" error, add `nlp.add_pipe('sentencizer')`.
    try:
        if nlp is not None:
            # If pipeline lacks components that set sentence boundaries, add sentencizer
            if 'sentencizer' not in getattr(nlp, 'pipe_names', []):
                try:
                    nlp.add_pipe('sentencizer')
                except Exception:
                    # some spacy versions may require a different name or component; ignore if it fails
                    pass
    except Exception:
        # Be resilient: if anything goes wrong here, fall back to None and use naive splitter
        nlp = None

    def get_sentences(text):
        if nlp is not None:
            return [s.text.strip() for s in nlp(text).sents]
        # naive fallback: split on common Chinese/Japanese/English sentence boundaries
        import re

        parts = re.split(r'[。！？\n]+', text)
        parts = [p.strip() for p in parts if p and p.strip()]
        if not parts:
            return [text.strip()]
        return parts

    results = []
    for item in items:
        # patch evaluate_item call to use get_sentences when nlp is None
        if nlp is None:
            # inline behavior of evaluate_item but with fallback splitter
            passage = item.get("passage", "")
            sampled = item.get("sampled_passages", [])
            sentences = get_sentences(passage)
            scores = checker.predict(
                sentences=sentences,
                passage=passage,
                sampled_passages=sampled,
                num_questions_per_sent=item.get("num_questions_per_sent", 3),
                beta1=item.get("beta1", 0.8),
                beta2=item.get("beta2", 0.8),
            )
            res = {
                "id": item.get("id"),
                "passage": passage,
                "sentences": sentences,
                "sampled_passages": sampled,
                "scores": scores,
            }
        else:
            res = evaluate_item(item, checker, nlp)
        # compute answer verdict using metadata
        verdict_info = evaluate_answer_verdict(item, res["passage"], res["sentences"], res.get("sampled_passages", []))
        res["verdict"] = verdict_info
        results.append(res)

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_path = args.output or f"results/g_eval_selfcheck_results_{ts}.json"
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    # Ensure results are JSON-serializable (convert tensors, numpy types, datetimes, etc.)
    safe_results = make_json_serializable(results)
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(safe_results, f, ensure_ascii=False, indent=2)

    print(f"Wrote {len(results)} result(s) to {out_path}")

    # Try to write a .docx summary, fallback to .txt
    docx_path = out_path.replace('.json', '.docx')
    txt_path = out_path.replace('.json', '.txt')
    try:
        from docx import Document

        doc = Document()
        doc.add_heading('g-eval SelfCheck Benchmark Results', level=1)
        for r in results:
            doc.add_heading(r.get('id', ''), level=2)
            meta = next((it.get('metadata') for it in items if it.get('id') == r.get('id')), {})
            if meta:
                q = meta.get('question') or meta.get('en')
                if q:
                    doc.add_paragraph(f"Question: {q}")
            doc.add_paragraph(f"Passage: {r.get('passage')}")
            doc.add_paragraph(f"Sentences: {r.get('sentences')}")
            doc.add_paragraph(f"Sampled: {r.get('sampled_passages')}")
            doc.add_paragraph(f"Scores: {r.get('scores')}")
            v = r.get('verdict', {})
            doc.add_paragraph(f"Verdict: {v.get('verdict')}")
            doc.add_paragraph(f"Found true labels: {v.get('found_true')}")
            doc.add_paragraph(f"Found false labels: {v.get('found_false')}")
            doc.add_paragraph('---')
        doc.save(docx_path)
        print(f"Also wrote docx summary to {docx_path}")
    except Exception:
        with open(txt_path, 'w', encoding='utf-8') as tf:
            tf.write('g-eval SelfCheck Benchmark Results\n')
            for r in results:
                tf.write(f"ID: {r.get('id')}\n")
                meta = next((it.get('metadata') for it in items if it.get('id') == r.get('id')), {})
                if meta:
                    q = meta.get('question') or meta.get('en')
                    if q:
                        tf.write(f"Question: {q}\n")
                tf.write(f"Passage: {r.get('passage')}\n")
                tf.write(f"Sentences: {r.get('sentences')}\n")
                tf.write(f"Sampled: {r.get('sampled_passages')}\n")
                tf.write(f"Scores: {r.get('scores')}\n")
                v = r.get('verdict', {})
                tf.write(f"Verdict: {v.get('verdict')}\n")
                tf.write(f"Found true labels: {v.get('found_true')}\n")
                tf.write(f"Found false labels: {v.get('found_false')}\n")
                tf.write('---\n')
        print(f"Also wrote text summary to {txt_path}")


if __name__ == "__main__":
    main()
